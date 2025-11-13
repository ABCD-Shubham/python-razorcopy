from fastapi import FastAPI, HTTPException, Request, Depends, Query, Body
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import text
from app.models.schemas import (
    RequestData,
    RequestData2,
    CompanyDetails,
    RequestData3,
    Sitemap,
    URLInput,
    MetaAnalysisResult,
    RequestDataSite,
    ArticleRequest,
    TitlesRequest,
    CompanyOutline,
    Titleoutline,
    PdfData,
    outline,
    FindTitle,
)
from app.services.google_search import extract_content
from app.services.scraper import (
    target_audience_generator,
    generate_preview,
    generate_target_audience,
    get_sitemap_urls,
    target_audience_generator1,
    generates_previews,
    article_google_search_links2,
    generate_owner_bio
)
from app.api.endpoints.company_overview import company_overview, company_overview1
from contextlib import asynccontextmanager
import os
from sqlalchemy import text
from dotenv import load_dotenv
import logging
import asyncio
from openai import OpenAI
from aiohttp import ClientSession
from app.services.url_analist import analyze_batch_urls
from app.services.sitemap_parser import SitemapParser
from typing import List
import google.generativeai as genai
import requests
from app.api.endpoints.company_business_summary import extract_content1
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.models.project_model import (
    Project,
    Article,
    Guideline,
    SystemPrompt,
    PromptTypes,
)
import json
import base64
from fastapi import APIRouter
from sqlalchemy.future import select
import httpx
import uuid
from fuzzywuzzy import fuzz
from fastapi import FastAPI, UploadFile, File
import fitz
from docx import Document
import io
import subprocess
import logging
import os
import tempfile
import zipfile
from typing import List, Optional
import re
import traceback
from uuid import UUID
import asyncio
import re
from typing import List, Optional
from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig
from playwright.async_api import async_playwright
import aiohttp
import anthropic
from urllib.parse import urlparse
from app.seo_audit.router import seo_audit_router
from app.seo_audit.models import Base
from app.core.database import engine


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
CONCURRENT_REQUESTS = 5
semaphore = asyncio.Semaphore(CONCURRENT_REQUESTS)
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
CUSTOM_GOOGLE_SEARCH = os.getenv("CUSTOM_GOOGLE_SEARCH")
CX_ID = os.getenv("CX_ID")

router = APIRouter()
api_key = CUSTOM_GOOGLE_SEARCH
cx = CX_ID

if not OPENAI_API_KEY:
    raise ValueError("Missing OpenAI API key.")
if not GEMINI_API_KEY:
    raise ValueError("Missing Gemini API key.")
if not CLAUDE_API_KEY:
    raise ValueError("Missing Claude API Key.")
if not CUSTOM_GOOGLE_SEARCH or not CX_ID:
    logger.warning(
        "Google Custom Search API keys are missing. Some features may not work."
    )

# claude client
claude_client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

# OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)
model = "ft:gpt-4.1-2024-08-06:e2m::ApEMBO4D"
genai.configure(api_key=GEMINI_API_KEY)
model_name = "gemini-1.5-pro-latest"


app = FastAPI()


origins = os.getenv("PY_PORT")

app.add_middleware(
    CORSMiddleware,
    allow_origins=(
        origins
        if origins
        else [
            "http://localhost:3001/",
            "https://app.razorcopy.com/",
            "https://rz-react.sitepreviews.dev/",
        ]
    ),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_URL_PREVIEW = os.getenv("BASE_URL_PREVIEW")
BASE_URL = os.getenv("BASE_URL")
WEBHOOK_AUTH_TOKEN = os.getenv("WEBHOOK_AUTH_TOKEN")

WEBHOOK_BASE_URL = os.getenv("WEBHOOK_BASE_URL", "http://192.168.1.75:8001/webhooks")
USER_AGENT = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.5481.100 Safari/537.36"
PLAYWRIGHT_TIMEOUT = 45000
MAX_RETRIES = 3
SEM_LIMIT = 5


@app.get("/health")
def health_check(db: Session = Depends(get_db)):
    try:
        # Execute a simple query to test the connection
        db.execute(text("SELECT 1")).scalar()
        return {"status": "healthy", "database": "connected"}
    except Exception as e:
        return {"status": "unhealthy", "database": "disconnected", "error": str(e)}


@app.post("/company-business-summary")
def scrape_company_details(request_data: RequestData2):
    try:
        company_name = request_data.company_name
        result = extract_content1(company_name, num=6)

        # Get the JSON response from the refactored function
        company_overview_json = company_overview1(result)

        # Extract individual fields from the JSON response
        company_details = company_overview_json.get("company_details", "")
        extracted_company_name = company_overview_json.get("company_name", company_name)
        owner_bio = company_overview_json.get("owner_bio", "")
        target_audience_string = company_overview_json.get("target_audience", "")
        
        # Convert comma-separated string to array
        target_audience = []
        if target_audience_string and target_audience_string != "Information not available":
            target_audience = [audience.strip() for audience in target_audience_string.split(",") if audience.strip()]

        # Format the company details (optional - you can remove this if you want to keep original formatting)
        formatted_details = company_details.replace(":", ":\n\n")

        # Return the structured response with all four fields
        return {
            "company_name": extracted_company_name,
            "company_details": formatted_details,
            "owner_bio": owner_bio,
            "target_audience": target_audience
        }

    except Exception:
        error_trace = traceback.format_exc()
        with open("error_log.txt", "w", encoding="utf-8") as f:
            f.write(error_trace)
        raise HTTPException(status_code=500, detail="Internal server error. See logs.")


@app.post("/target-audience")
def get_target_audience(request_data: CompanyDetails):
    company_details = request_data.company_details
    target_audience = generate_target_audience(company_details)
    return target_audience


@app.post("/generate-outline")
def generates_contents(
    request_data: outline, db: Session = Depends(get_db), article_id=None
):
    try:
        logger.info(
            f"Received /generate-outline request for articleId: {request_data.articleId}"
        )

        try:
            article_uuid = uuid.UUID(request_data.articleId)
        except ValueError:
            logger.warning(
                f"Invalid UUID format for articleId: {request_data.articleId}"
            )
            raise HTTPException(
                status_code=400, detail="Invalid UUID format for articleId"
            )

        # Fetch the target article
        article_query = select(Article).where(Article.id == article_uuid)
        article_result = db.execute(article_query).unique()
        article = article_result.scalars().first()

        if not article:
            logger.warning(f"Article not found for articleId: {request_data.articleId}")
            raise HTTPException(status_code=404, detail="Article not found")

        # Fetch associated project
        if not article.projectId:
            logger.error(
                f"Article with ID {request_data.articleId} has no associated project"
            )
            raise HTTPException(
                status_code=400, detail="Article has no associated project"
            )

        try:
            project_uuid = uuid.UUID(str(article.projectId))
        except ValueError:
            logger.error(f"Invalid UUID format for projectId: {article.projectId}")
            raise HTTPException(
                status_code=400, detail="Invalid UUID format for projectId"
            )

        project_query = select(Project).where(Project.id == project_uuid)
        project_result = db.execute(project_query).unique()
        project = project_result.scalars().first()

        if not project:
            logger.warning(f"Project not found for projectId: {article.projectId}")
            raise HTTPException(status_code=404, detail="Project not found")

        articles = [
            {
                "id": article.id,
                "secondary_keywords": article.secondary_keywords,
                "name": article.name,
                "keywords": article.keywords,
            }
        ]

        logger.info(
            f"Generating content preview for Article: {article.name}, Project: {project.id}"
        )

        client_site_url = project.description
        keywords = [a["keywords"] for a in articles]
        title = [a["name"] for a in articles]
        secondary_keywords = [a["secondary_keywords"] for a in articles]
        target_audience = project.targeted_audience

        # Generate preview content
        generated_content_preview = generates_previews(
            title,
            keywords,
            target_audience,
            secondary_keywords,
            client_site_url,
            article,
            db,
        )

        logger.info(
            f"Successfully generated content preview for articleId: {article.id}"
        )

        return generated_content_preview

    except Exception as e:
        logger.exception(
            f"Unhandled error in /generate-outline for articleId: {request_data.articleId}"
        )
        raise HTTPException(status_code=500, detail=str(e))


async def count_urls_in_sitemaps(sitemap_urls):
    """Count URLs in sitemaps."""
    total_count = 0
    details = []
    for sitemap_url in sitemap_urls:
        urls = get_sitemap_urls(sitemap_url)
        count = len(urls)
        match = re.search(r"/([^/]+?)-sitemap", sitemap_url)
        if match:
            sitemap_type = match.group(1)
            details.append({"type": sitemap_type, "count": count})
        total_count += count
    return total_count, details


async def determine_sitemap(user_site):
    sitemap_index = user_site.rstrip("/") + "/sitemap_index.xml"
    sitemap = user_site.rstrip("/") + "/sitemap.xml"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.5735.90 Safari/537.36"
    }
    async with ClientSession(headers=headers) as session:
        try:
            async with session.get(sitemap_index, timeout=20) as response:
                if response.status == 200:
                    return sitemap_index
        except:
            pass
        try:
            async with session.get(sitemap, timeout=20) as response:
                if response.status == 200:
                    return sitemap
        except:
            pass
    return None


@app.post("/fetch-sitemaps")
async def fetch_sitemaps(request_data: Sitemap):
    try:
        company_name = request_data.company_name
        logger.info(f"Fetching sitemaps for company: {company_name}")

        result = extract_content(company_name, num=6)
        user_site = request_data.company_name
        sitemap_url = await determine_sitemap(user_site)

        if not sitemap_url:
            logger.error("No sitemap found for the given URL")
            raise HTTPException(status_code=404, detail="No sitemap found")

        sitemap_urls = get_sitemap_urls(sitemap_url)
        total_pages, details = await count_urls_in_sitemaps(sitemap_urls)

        logger.info(f"Sitemap fetched successfully: {sitemap_url}")
        return {
            "site_urls": sitemap_url,
            "total_pages": total_pages,
            "contain_types": details,
        }

    except Exception as e:
        logger.error(f"An error occurred: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Internal Server Error")


async def process_urls_in_batches(urls, batch_size=50):
    """
    Processes URLs in smaller batches to prevent memory exhaustion.
    """
    results = []
    for i in range(0, len(urls), batch_size):
        batch = urls[i : i + batch_size]
        async with semaphore:
            logger.info(
                f"Processing batch {i//batch_size + 1}/{(len(urls)//batch_size) + 1}"
            )
            batch_results = await analyze_batch_urls(batch)
            results.extend(batch_results)

        await asyncio.sleep(0.1)

    return results


@app.post("/sitemap", response_model=List[MetaAnalysisResult])
async def fetch_and_analyze_sitemap(url_input: URLInput):
    """
    Fetch all URLs from a sitemap and analyze them in batches.
    This replaces `discover-urls` and `analyze-batch` by combining their logic.
    """
    try:
        logger.info(f"Fetching URLs from sitemap: {url_input.url}")
        async with semaphore:
            parser = SitemapParser(str(url_input.url))

            try:
                urls = await asyncio.wait_for(parser.get_all_urls(), timeout=300)
                content_types = parser.content_types
            except asyncio.TimeoutError:
                logger.warning("URL discovery timed out, returning partial results")
                urls = list(parser.discovered_urls)
                content_types = parser.content_types

        logger.info(f"Discovered {len(urls)} URLs from sitemap.")

        results = await process_urls_in_batches(urls, batch_size=100)

        for result in results:
            result.content_type = content_types.get(result.url, "Unknown")

        logger.info(f"Batch analysis completed for {len(results)} URLs.")
        return results

    except Exception as e:
        logger.error(f"Error processing sitemap: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.post("/get-titles")
def get_all_titles(request: TitlesRequest, db: Session = Depends(get_db)):
    try:
        # Validate ProjectId format
        try:
            project_uuid = uuid.UUID(request.ProjectId)
        except ValueError:
            raise HTTPException(
                status_code=400, detail="Invalid UUID format for ProjectId"
            )

        # Fetch project details
        project_query = select(Project).where(Project.id == project_uuid)
        project_result = db.execute(project_query).unique()
        project = project_result.scalars().first()

        if not project:
            logger.error(f"Project not found for ID: {request.ProjectId}")
            raise HTTPException(status_code=404, detail="Project not found")

        # Decode detailed sitemap if available
        detailed_sitemap = None
        if project.detailedsitemap and project.detailedsitemap.strip():
            try:
                decoded_sitemap = decode_sitemap(project.detailedsitemap)
                if isinstance(decoded_sitemap, dict):
                    detailed_sitemap = decoded_sitemap
                else:
                    logger.error(
                        f"Decoded detailedsitemap is not a valid dictionary: {decoded_sitemap}"
                    )
            except Exception as e:
                logger.error(
                    f"Failed to decode detailedsitemap JSON: {e} - Raw Data: {project.detailedsitemap}"
                )

        # Fetch articles associated with the project
        article_query = select(Article).where(Article.projectId == project_uuid)
        articles = db.execute(article_query).scalars().all()

        if not articles:
            logger.warning(f"No articles found for project ID: {request.ProjectId}")

        # Extract keywords from request
        prompt_type_ids = [item.promptTypeId for item in request.Keywords]
        keywords = [item.keyword for item in request.Keywords]
        logger.info(f"Extracted {len(keywords)} keywords from request: {keywords}")

        # Fetch PromptTypes and map promptTypeId → topic_prompt_id
        prompt_types = (
            db.execute(select(PromptTypes).where(PromptTypes.id.in_(prompt_type_ids)))
            .scalars()
            .all()
        )
        prompt_type_map = {pt.id: pt.topic_prompt_id for pt in prompt_types}

        # Fetch SystemPrompts using topic_prompt_ids
        topic_prompt_ids = list(set(filter(None, prompt_type_map.values())))
        system_prompts = (
            db.execute(
                select(SystemPrompt).where(SystemPrompt.id.in_(topic_prompt_ids))
            )
            .scalars()
            .all()
        )
        system_prompt_map = {sp.id: sp.description for sp in system_prompts}

        # Pair each keyword with its system prompt description
        keyword_prompt_pairs = []
        for item in request.Keywords:
            prompt_type_id = UUID(item.promptTypeId)
            keyword = item.keyword
            topic_prompt_id = prompt_type_map.get(prompt_type_id)
            prompt_description = system_prompt_map.get(
                topic_prompt_id, "No prompt found"
            )

            keyword_prompt_pairs.append(
                {
                    "keyword": keyword,
                    "promptTypeId": str(prompt_type_id),
                    "topic_prompt_id": (
                        str(topic_prompt_id) if topic_prompt_id else None
                    ),
                    "prompt_description": prompt_description,
                }
            )

        logger.info(f"Keyword-Prompt pairs: {keyword_prompt_pairs}")
        final_titles = []

        # Define titles information
        titles_info = {
            "id": str(project.id),
            "name": project.name,
            # "keywords": [pair["keyword"] for pair in keyword_prompt_pairs],
            # "system_prompt": [pair["prompt_description"] for pair in keyword_prompt_pairs],
            "language": project.language,
            "website_url": project.website_url,
            "competitors_websites": project.competitors_websites,
            "location": project.location,
            "targeted_audience": project.targeted_audience,
            # "created_at": project.created_at,
            # "updated_at": project.updated_at
        }
        # print(titles_info,"----------")

        # Function to generate prompt
        def generate_prompt(keyword, system_prompt_description, titles_info):
            project_name = titles_info.get("name")
            project_target_audience = titles_info.get("targeted_audience")
            project_website_url = titles_info.get("website_url")
            project_competitors_websites = titles_info.get("competitors_websites")
            project_language = titles_info.get("language")
            project_location = titles_info.get("location")
            # secondary_keywords = ", ".join(titles_info["secondary_keywords"])
            # keywords = ", ".join(titles_info.get('keywords', []))

            prompt = system_prompt_description.format(
                project_name=project_name,
                project_website_url=project_website_url,
                project_target_audience=project_target_audience,
                project_competitors_websites=project_competitors_websites,
                project_language=project_language,
                project_location=project_location,
                # keywords=", ".join(titles_info.get('keywords', [])),
                keywords=keyword,
            )
            return prompt

        # Loop through keyword-prompt pairs and generate SEO titles
        for pair in keyword_prompt_pairs:
            keyword = pair["keyword"]
            system_prompt_description = pair["prompt_description"]

            prompt = generate_prompt(keyword, system_prompt_description, titles_info)

            # Request OpenAI API for SEO-optimized blog titles
            messages = [
                {"role": "system", "content": prompt},
                {"role": "user", "content": "Generate SEO-optimized blog titles."},
            ]

            try:
                response = requests.post(
                    "https://api.openai.com/v1/chat/completions",
                    json={"model": "gpt-4.1", "messages": messages, "temperature": 0.8},
                    headers={
                        "Authorization": f"Bearer {OPENAI_API_KEY}",
                        "Content-Type": "application/json",
                    },
                )

                response_data = response.json()
                generated_titles = (
                    response_data.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                    .strip()
                    .split("\n")
                )

                # Clean the generated titles
                cleaned_titles = [
                    re.sub(r"^[-\d\.]+\s*", "", title).strip()
                    for title in generated_titles
                    if title.strip()
                ]

                # Check for title similarity
                matched_titles = []
                for article in articles:
                    original_title = article.name.strip() if article.name else ""

                    for gen_title in cleaned_titles:
                        similarity_original = fuzz.ratio(
                            original_title.lower(), gen_title.lower()
                        )

                        if similarity_original >= 80:
                            logger.info(
                                f"Generated title '{gen_title}' is too similar to existing titles."
                            )
                            unique_title_prompt = f"Generate a completely unique blog title that is different from '{gen_title}'."

                            try:
                                unique_response = requests.post(
                                    "https://api.openai.com/v1/chat/completions",
                                    json={
                                        "model": "gpt-4.1",
                                        "messages": [
                                            {
                                                "role": "user",
                                                "content": unique_title_prompt,
                                            }
                                        ],
                                        "temperature": 0.8,
                                    },
                                    headers={
                                        "Authorization": f"Bearer {OPENAI_API_KEY}",
                                        "Content-Type": "application/json",
                                    },
                                )

                                unique_data = unique_response.json()
                                new_title = (
                                    unique_data.get("choices", [{}])[0]
                                    .get("message", {})
                                    .get("content", "")
                                    .strip()
                                )
                                # matched_titles.append({"original": gen_title, "unique": new_title})

                            except requests.exceptions.RequestException as e:
                                logger.error(f"Error generating unique title: {e}")

                # Add unique titles or cleaned titles to final list
                if matched_titles:
                    final_titles.append(matched_titles[0])
                elif cleaned_titles:
                    final_titles.append(cleaned_titles[0])
                else:
                    final_titles.append("No title generated")
                # final_titles.append({
                #     "keyword": keyword,
                #     "titles": final_title
                # })
                # final_titles.append(final_title)

            except requests.exceptions.RequestException as e:
                logger.error(f"OpenAI API request error: {e}")
                raise HTTPException(
                    status_code=500,
                    detail="Failed to generate titles due to API error.",
                )

        return final_titles

    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        raise HTTPException(status_code=500, detail="An unexpected error occurred.")


@app.post("/check-title")
def check_title(request: FindTitle, db: Session = Depends(get_db)):
    title_to_check = request.title
    projectId = request.ProjectId

    if not title_to_check:
        return {"error": "Title is required to check"}

    exists = (
        db.query(Article)
        .filter(Article.name == title_to_check, Article.projectId == projectId)
        .first()
    )

    if exists:
        return exists.name
    else:
        return None


def calculate_similarity_percentage(text1: str, text2: str) -> float:
    """Calculate a simple similarity percentage between two strings."""
    set1, set2 = set(text1.lower().split()), set(text2.lower().split())
    intersection = set1.intersection(set2)
    return (len(intersection) / max(len(set1), len(set2))) * 100 if set1 and set2 else 0


def decode_sitemap(sitemap_data: str):
    if not sitemap_data:
        return {}

    try:
        json_bytes = base64.b64decode(sitemap_data)
        json_str = json_bytes.decode("utf-8")
        data = json.loads(json_str)
    except (base64.binascii.Error, UnicodeDecodeError, json.JSONDecodeError):
        try:
            data = json.loads(sitemap_data)
        except json.JSONDecodeError:
            logger.error("Failed to decode sitemap data")
            return {}

    if isinstance(data, list):
        try:
            return {str(i): item for i, item in enumerate(data)}
        except Exception:
            logger.error("Failed to convert list to dictionary")
            return {}

    if not isinstance(data, dict):
        logger.error(f"Decoded sitemap is not a dictionary: {data}")
        return {}
    return data


def count_words(text):
    return len(text.split())


def calculate_average_word_count(word_counts):
    return sum(word_counts) / len(word_counts) if word_counts else 0


def extract_citations(content: str) -> List[str]:
    matches = re.findall(r'Source:\s*(https?://[^\s,")]+)', content)
    all_urls = re.findall(r'https?://[^\s,")]+', content)
    citations = matches + all_urls
    # Remove duplicates while preserving order
    seen = set()
    unique_citations = []
    for url in citations:
        if url not in seen:
            seen.add(url)
            unique_citations.append(url)
    return unique_citations


async def scrape_retries(
    query: str, api_key: str, cx: str, num_results: int = 5
) -> List[str]:
    """
    Fetches search result URLs from the Google Custom Search API.
    Returns a list of result URLs.
    """
    search_url = "https://www.googleapis.com/customsearch/v1"
    params = {"key": api_key, "cx": cx, "q": query, "num": num_results}

    async with aiohttp.ClientSession() as session:
        async with session.get(search_url, params=params) as response:
            if response.status != 200:
                error_text = await response.text()
                raise Exception(f"Google API error: {response.status} - {error_text}")

            data = await response.json()
            if "items" not in data:
                return []

            return [item.get("link") for item in data["items"]]


async def fetch_search_results(
    query: str, api_key: str, cx: str, num_results: int = 5
) -> List[str]:
    search_url = "https://www.googleapis.com/customsearch/v1"
    params = {"key": api_key, "cx": cx, "q": query, "num": num_results}

    async with aiohttp.ClientSession() as session:
        async with session.get(search_url, params=params) as response:
            if response.status != 200:
                error_text = await response.text()
                raise Exception(f"Google API error: {response.status} - {error_text}")
            data = await response.json()
            return [item.get("link") for item in data.get("items", [])]


async def extract_content_google(query: str, api_key: str, cx: str, num: int = 5):
    # Fetch clean URLs via Google Custom Search
    links = await fetch_search_results(query, CUSTOM_GOOGLE_SEARCH, cx, num_results=5)

    # Filter out image URLs and ensure we have valid article URLs
    valid_links = []
    for link in links:
        # Skip image files and other non-article URLs
        if not any(
            ext in link.lower()
            for ext in [".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg"]
        ):
            valid_links.append(link)

    links = valid_links

    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.5481.100 Safari/537.36"
    }
    semaphore = asyncio.Semaphore(5)

    async def scrape_with_retries(url: str, max_attempts: int = 3) -> Optional[str]:
        for attempt in range(max_attempts):
            try:
                async with semaphore:
                    async with async_playwright() as p:
                        browser = await p.chromium.launch(
                            headless=True,
                            args=[
                                "--disable-gpu",
                                "--no-sandbox",
                                "--disable-dev-shm-usage",
                            ],
                        )
                        context = await browser.new_context(
                            user_agent=headers["User-Agent"],
                            viewport={"width": 1920, "height": 1080},
                        )

                        # Enable request interception to block unnecessary resources
                        await context.route(
                            "**/*.{png,jpg,jpeg,gif,svg,css,font,woff,woff2,eot,ttf,otf}",
                            lambda route: route.abort(),
                        )

                        page = await context.new_page()

                        # Set longer timeout for initial page load
                        await page.goto(
                            url, timeout=60000, wait_until="domcontentloaded"
                        )

                        # Wait for main content to be available
                        await page.wait_for_selector("body", timeout=10000)

                        # Extract main content
                        content = await page.evaluate(
                            """() => {
                            // Remove unwanted elements
                            const elementsToRemove = document.querySelectorAll(
                                'header, footer, nav, script, style, iframe, .ad, .advertisement, ' +
                                '.banner, .popup, .modal, .cookie-banner, .newsletter, .social-share'
                            );
                            elementsToRemove.forEach(el => el.remove());
                            
                            // Get main content
                            const mainContent = document.querySelector('main, article, .content, .post, .article') || document.body;
                            return mainContent.innerText;
                        }"""
                        )

                        await browser.close()

                        if content and len(content.strip()) > 0:
                            # Clean up the content
                            cleaned_content = " ".join(content.split())
                            footer = f"\n\nSource: {url}\n"
                            return f"{cleaned_content}{footer}"
                        else:
                            raise Exception("No content extracted")

            except Exception as e:
                print(
                    f"[Attempt {attempt+1}] Failed scraping {url}: {type(e).__name__}: {str(e)[:100]}..."
                )
                await asyncio.sleep(2**attempt)
        return None

    # Scrape pages concurrently
    tasks = [scrape_with_retries(url) for url in links]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    # Filter successful scrapes and get their URLs
    successful_urls = []
    related_pages = []
    word_counts = []
    failed_links = []
    citations_text = ""

    for link, content in zip(links, results):
        if isinstance(content, Exception) or content is None:
            failed_links.append(link)
            continue

        successful_urls.append(link)
        wc = count_words(content)
        word_counts.append(wc)
        related_pages.append((link, content, wc))
        citations_text += content + "\n"

    # Take only the first 5 successful URLs
    successful_urls = successful_urls[:num]

    # Print successful URLs
    print("\nTop 5 Successfully Scraped URLs:")
    for i, url in enumerate(successful_urls, 1):
        print(f"{i}. {url}")
    print("\n")

    if failed_links:
        print("\nFailed to scrape:")
        for link in failed_links:
            print(f"- {link}")

    # Prepare result text
    result_text = "Reference Articles:\n"
    for idx, (url, content, wc) in enumerate(related_pages[:num], 1):
        result_text += f"{idx}. {url}\n"
        result_text += f"Article Content (Truncated):\n{content[:20000]}...\n\n"

    avg_word_count = calculate_average_word_count(word_counts)

    return result_text, avg_word_count, successful_urls


@app.post("/get-articles")
async def get_all_projects(request: ArticleRequest, db: Session = Depends(get_db)):
    """
    Retrieve all projects from the database.
    """
    try:
        try:
            article_uuid = uuid.UUID(request.articleId)
        except ValueError:
            raise HTTPException(
                status_code=400, detail="Invalid UUID format for articleId"
            )

        article_query = select(Article).where(Article.id == article_uuid)
        article_result = db.execute(article_query).unique()
        article = article_result.scalars().first()

        if not article:
            logger.error(f"Article not found for ID: {request.articleId}")
            return {
                "message": "Article not found",
                "article": None,
                "project": None,
                "guideline": None,
                "system_prompts": None,
            }

        project = None
        if article.projectId:
            try:
                project_uuid = uuid.UUID(str(article.projectId))
                project_query = select(Project).where(Project.id == project_uuid)
                project_result = db.execute(project_query).unique()
                project = project_result.scalars().first()
            except ValueError:
                logger.error(f"Invalid UUID format for project_id: {article.projectId}")

        guideline = None
        if project and project.guideline_id:
            try:
                guideline_uuid = uuid.UUID(str(project.guideline_id))
                guideline_query = select(Guideline).where(
                    Guideline.id == guideline_uuid
                )
                guideline_result = db.execute(guideline_query).unique()
                guideline = guideline_result.scalars().first()
            except ValueError:
                logger.error(
                    f"Invalid UUID format for guideline_id: {project.guideline_id}"
                )

        if article.prompt_type_id:
            prompt_type = (
                db.execute(
                    select(PromptTypes).where(PromptTypes.id == article.prompt_type_id)
                )
                .scalars()
                .first()
            )

        if prompt_type and prompt_type.article_prompt_id:
            system_prompt = (
                db.execute(
                    select(SystemPrompt).where(
                        SystemPrompt.id == prompt_type.article_prompt_id,
                        SystemPrompt.type == "article",
                    )
                )
                .scalars()
                .first()
            )

        if system_prompt:
            logger.info(f"System prompt found: {system_prompt.description} (added)")
        else:
            logger.warning("No system prompt found. --------")

        result_text, avg_word_count, reference_links = await extract_content_google(
            article.name, api_key=CUSTOM_GOOGLE_SEARCH, cx=CX_ID, num=5
        )

        # Format top 5 URLs for response
        formatted_top_urls = "\n".join(
            [f"{i+1}. {url}" for i, url in enumerate(reference_links)]
        )

        logger.info(f"Successfully scraped {len(reference_links)} URLs")

        # Prepare response data
        article_info = {
            "id": article.id,
            "name": article.name,
            "system_prompt": {"description": system_prompt.description},
            "generate_outline": article.generated_outline,
            "article_rulesets": article.article_rulesets,
            "scraped_content": result_text,
            "reference_links": formatted_top_urls,
            "avg_word_count": avg_word_count,
            "project": (
                {
                    "id": str(project.id) if project else None,
                    "name": project.name if project else None,
                    "description": project.description if project else None,
                    "language": project.language if project else None,
                    "location": project.location if project else None,
                    "targeted_audience": project.targeted_audience if project else None,
                    "created_at": project.created_at if project else None,
                    "guideline_description": (
                        project.guideline_description if project else None
                    ),
                    "updated_at": project.updated_at if project else None,
                    "organization_archetype": (
                        project.organization_archetype
                        if project and project.organization_archetype
                        else None
                    ),
                    "brand_spokesperson": (
                        project.brand_spokesperson
                        if project and project.brand_spokesperson
                        else None
                    ),
                    "most_important_thing": (
                        project.most_important_thing
                        if project and project.most_important_thing
                        else None
                    ),
                    "unique_differentiator": (
                        project.unique_differentiator
                        if project and project.unique_differentiator
                        else None
                    ),
                    "author_bio": (
                        project.author_bio if project and project.author_bio else None
                    ),
                    "guideline": (
                        {"description": guideline.description if guideline else None}
                        if guideline
                        else None
                    ),
                }
                if project
                else None
            ),
        }

        def generate_prompt(article_info, formatted_references_prompt):
            """
            Create a structured prompt for AI summarization using the database content.
            """
            article = article_info
            scraped_content = article.get("scraped_content", "")
            truncated_content = truncate_text(scraped_content)
            project = article.get("project", {})
            project_name = project.get("name", "N/A")
            project_description = project.get("description", "N/A")
            project_language = project.get("language", "N/A")
            project_location = project.get("location", "N/A")
            project_target_audience = project.get("targeted_audience", "N/A")
            project_industry_description = project.get("guideline", {}).get(
                "description", "N/A"
            )
            project_general_guideline = project.get("guideline_description", "N/A")

            # Get additional fields from project
            organization_archetype = project.get("organization_archetype", "N/A")
            brand_spokesperson = project.get("brand_spokesperson", "N/A")
            most_important_thing = project.get("most_important_thing", "N/A")
            unique_differentiator = project.get("unique_differentiator", "N/A")
            author_bio = project.get("author_bio", "N/A")

            articles = article.get("name")
            project_outline = article.get("generate_outline")
            article_rulesets = article.get("article_rulesets", "")
            formatted_references_prompt = "\n".join(reference_links)

            formatted_content = {
                "truncated_content": truncated_content,
                "formatted_references": formatted_references_prompt,
                "project": article.get("project", {}),
                "project_name": project_name,
                "article": articles,
                "project_target_audience": project_target_audience,
                "project_description": project_description,
                "project_language": project_language,
                "project_location": project_location,
                "project_outline": project_outline,
                "project_general_guideline": project_industry_description,
                "project_industry_description": project_general_guideline,
                "organization_archetype": organization_archetype,
                "brand_spokesperson": brand_spokesperson,
                "most_important_thing": most_important_thing,
                "unique_differentiator": unique_differentiator,
                "author_bio": author_bio,
                "article_rulesets": article_rulesets,
            }

            prompt_template = article_info.get("system_prompt", {}).get(
                "description", ""
            )
            
            # Add article rulesets to the prompt if they exist
            article_rulesets_section = ""
            if article_rulesets and article_rulesets.strip():
                article_rulesets_section = f"\n\n**Article-Specific Rules and Guidelines:**\n{article_rulesets}\n\nPlease follow these specific rules and guidelines when generating the content for this article."
                logger.info(f"Article rulesets found and added to prompt: {article_rulesets[:100]}...")
            else:
                logger.info("No article rulesets found or rulesets are empty")
            
            # Add default content rules to prevent extra explanatory notes
            default_content_rules = """
**Strict Content Rules:**
- Write ONLY the article content
- NO explanatory notes, disclaimers, or meta-commentary
- NO language style notes or writing methodology explanations
- NO phrases like "Note on Language Style:" or similar
- Deliver pure article content only"""
            
            final_prompt = (
                prompt_template.format(**formatted_content)
                + article_rulesets_section
                + default_content_rules
                + "\n\n"
                + formatted_content["project_industry_description"]
                + formatted_content["formatted_references"]
            )
            
            # Log the final prompt length and if it contains article rulesets
            logger.info(f"Final prompt length: {len(final_prompt)} characters")
            if "Article-Specific Rules and Guidelines:" in final_prompt:
                logger.info("✅ Article rulesets successfully included in final prompt")
            else:
                logger.info("❌ Article rulesets NOT found in final prompt")
            
            return final_prompt

        prompt = generate_prompt(article_info, reference_links)
        summaries = {}

        if request.model in [None, ""]:
            summaries["open_ai"] = get_openai_summary(prompt, formatted_references)
            summaries["gemini"] = get_gemini_summary(prompt, formatted_references)
            summaries["claude"] = get_claude_summary(prompt, formatted_references)

        elif request.model == "open_ai":
            summaries["open_ai"] = get_openai_summary(prompt, reference_links)

        elif request.model == "gemini":
            summaries["gemini"] = get_gemini_summary(prompt, reference_links)

        elif request.model == "claude":
            summaries["claude"] = get_claude_summary(prompt, reference_links)

        elif request.model == "claude":
            summaries["claude"] = get_claude_summary(prompt, formatted_references)

        else:
            raise HTTPException(status_code=400, detail="Invalid model specified")
        webhook_responses = {}
        word_counts = []

        for model_name, summary in summaries.items():
            word_count = len(summary.split())
            webhook_url = f"{BASE_URL}/webhooks/{request.articleId}/content"
            payload = {
                "model": model_name,
                "content": summary,
                "requestId": request.requestId,
                "avg_word_count": word_count,
            }
            logger.info(f"Average word count: {avg_word_count}")

            try:
                with httpx.Client(timeout=10) as client:
                    response = client.post(
                        webhook_url,
                        json=payload,
                        headers={"Authorization": f"Bearer {WEBHOOK_AUTH_TOKEN}"},
                    )

                if response.status_code not in [200, 202]:
                    logger.error(
                        f"Webhook failed for {model_name}: {response.status_code} - {response.text}"
                    )
                    webhook_responses[model_name] = summary
                else:
                    webhook_responses[model_name] = summary

            except Exception as e:
                logger.error(f"Error sending {model_name} webhook: {e}")
                webhook_responses[model_name] = summary

        return {"webhook_responses": webhook_responses, "avg_word_count": word_count}

    except Exception as e:
        logger.error(f"Error retrieving projects: {e}")
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


def truncate_text(text, max_words=500):
    """
    Convert input to string (if it's a tuple or list) and truncate to a specified word limit.
    """
    if isinstance(text, (tuple, list)):
        text = " ".join(map(str, text))

    text = str(text)
    words = text.split()[:max_words]
    return " ".join(words)


def get_claude_summary(prompt: str, formatted_references: list = None) -> str:
    try:
        response = claude_client.messages.create(
            model="claude-3-5-sonnet-20241022",  # or other available modelAdd commentMore actions
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
            system="You are a helpful assistant.",
        )
        output = (
            response.content[0].text
            if hasattr(response, "content")
            else response.get("completion", "")
        )
        # Append reference links, if provided
        if formatted_references:
            output += "\n\nReferences\n"
            output += "\n".join(
                [f"{i+1}. [{url}]({url})" for i, url in enumerate(formatted_references)]
            )
        return output
    except Exception as e:
        return f"Claude Error: {str(e)}"


def get_gemini_summary(prompt: str, formatted_references: List[str]):
    """
    Call Gemini AI to summarize the article and embed reference URLs at the end.
    Removes any AI-generated opening lines like 'Okay, here's a blog article...'
    """
    try:
        # Prepare prompt
        formatted_prompt = (
            "\n".join([item["content"] for item in prompt])
            if isinstance(prompt, list)
            else prompt
        )
        logger.debug(f"Formatted References for Gemini: {formatted_references}")

        # Initialize Gemini model
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(formatted_prompt)

        # Extract generated text
        output = response.text if response else "Error: No response from Gemini AI."
        word_count = len(output.split())

        print(f"Generated Word Count: {word_count}")
        if word_count < 2000:
            print(
                "Warning: The generated article is shorter than expected. Consider adjusting the prompt or regenerating."
            )

        # Remove the specific intro line if present
        intro_pattern = r"^Okay, here's a blog article.*?(?=\n\n|$)"
        output, num_subs = re.subn(
            intro_pattern, "", output, flags=re.IGNORECASE | re.DOTALL
        )
        if num_subs:
            print("Removed AI-generated intro line.")

        # Clean up any leading newlines
        output = output.lstrip()

        # Add top 5 URLs at the beginning of the content
        if formatted_references:
            reference_links = formatted_references[:5]  # Get top 5 URLs
            urls_section = "\nTop 5 Google Search URLs:\n"
            urls_section += "\n".join(
                [f"{i+1}. {url}" for i, url in enumerate(reference_links)]
            )
            output = output

            # Add references section at the end
            references_to_append = "\n\nReferences\n" + "\n".join(
                [f"{i+1}. [{url}]({url})" for i, url in enumerate(formatted_references)]
            )
            output += references_to_append
            logger.debug(f"References Appended:\n{references_to_append}")
        else:
            logger.debug(
                "No formatted references provided, skipping references section."
            )

        return output

    except Exception as e:
        print(f"Error generating content: {e}")
        return "Error generating content."


def get_openai_summary(prompt: str, citations: list = None) -> str:
    """
    Call OpenAI (GPT-4) to summarize the article and append reference URLs at the end.
    """
    try:
        completion = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are an AI assistant."},
                {"role": "user", "content": prompt},
            ],
        )
        output = completion.choices[0].message.content

        # Append reference links, if provided
        if citations:
            output += "\n\nReferences\n"
            output += "\n".join(
                [f"{i+1}. [{url}]({url})" for i, url in enumerate(citations)]
            )

        return output
    except Exception as e:
        return f"OpenAI Error: {str(e)}"


def get_claude_summary(prompt: str, formatted_references: list = None) -> str:
    try:
        response = claude_client.messages.create(
            model="claude-3-5-sonnet-20241022",  # or other available model
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
            system="You are a helpful assistant.",
        )
        output = (
            response.content[0].text
            if hasattr(response, "content")
            else response.get("completion", "")
        )
        # Append reference links, if provided
        if formatted_references:
            output += "\n\nReferences\n"
            output += "\n".join(
                [f"{i+1}. [{url}]({url})" for i, url in enumerate(formatted_references)]
            )
        return output
    except Exception as e:
        return f"Claude Error: {str(e)}"


def convert_doc_to_docx(doc_path: str) -> str:
    try:
        output_dir = os.path.dirname(doc_path)

        libreoffice_cmd = subprocess.run(
            ["which", "libreoffice"], capture_output=True, text=True
        ).stdout.strip()
        if not libreoffice_cmd:
            libreoffice_cmd = subprocess.run(
                ["which", "soffice"], capture_output=True, text=True
            ).stdout.strip()
        if not libreoffice_cmd:
            raise FileNotFoundError("LibreOffice not found. Please install it.")

        result = subprocess.run(
            [
                libreoffice_cmd,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                output_dir,
                doc_path,
            ],
            capture_output=True,
            text=True,
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

        new_doc_path = os.path.splitext(doc_path)[0] + ".docx"

        if not os.path.exists(new_doc_path):
            raise FileNotFoundError("Conversion failed: .docx file not created.")

        if not zipfile.is_zipfile(new_doc_path):
            raise ValueError("Converted .docx is not a valid zip archive.")

        return new_doc_path

    except Exception as e:
        raise RuntimeError(f"Error converting .doc to .docx: {e}")


@app.post("/file-ocr")
async def upload_file(file: UploadFile = File(...)):
    file_ext = file.filename.split(".")[-1].lower()

    if file_ext not in ["pdf", "docx", "doc"]:
        raise HTTPException(
            status_code=400, detail="Only PDF, DOCX, and DOC files are allowed."
        )

    try:
        file_data = await file.read()
        extracted_text = ""

        if file_ext == "pdf":
            pdf_document = fitz.open(stream=file_data, filetype="pdf")
            extracted_text = "\n".join([page.get_text("text") for page in pdf_document])

        elif file_ext == "docx":
            try:
                if not zipfile.is_zipfile(io.BytesIO(file_data)):
                    raise ValueError("Uploaded .docx file is not a valid archive.")
                doc = Document(io.BytesIO(file_data))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                raise HTTPException(
                    status_code=500, detail=f"Invalid .docx file: {str(e)}"
                )
        elif file_ext == "doc":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                temp_file.write(file_data)
                temp_path = temp_file.name
            try:
                converted_path = convert_doc_to_docx(temp_path)
                if not zipfile.is_zipfile(converted_path):
                    raise ValueError("Converted file is not a valid .docx archive.")
                with open(converted_path, "rb") as f:
                    doc = Document(f)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])
                os.remove(converted_path)
            except Exception as e:
                raise HTTPException(
                    status_code=500, detail=f"Error processing .doc file: {str(e)}"
                )
            finally:
                os.remove(temp_path)
        return {"filename": file.filename, "extracted_text": extracted_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")


app.include_router(seo_audit_router, prefix="/seo-audit", tags=["SEO Audit"])



@app.post("/owner-bio")
async def get_owner_bio(request_data: CompanyDetails):
    company_details = request_data.company_details
    owner_bio = generate_owner_bio(company_details)
    return owner_bio


@app.post("/generate-brand-spokesperson")
def generate_brand_spokesperson(request_data: dict):
    try:
        description = request_data.get("description", "")
        targeted_audience = request_data.get("targeted_audience", "")
        author_bio = request_data.get("author_bio", "")

        if not description or not targeted_audience or not author_bio:
            raise HTTPException(
                status_code=400, 
                detail="Description, targeted_audience, and author_bio are required"
            )

        # Create a comprehensive prompt for brand spokesperson generation
        prompt = f"""
Based on the following business information, suggest 3 famous people (real or fictional) who would be ideal spokespersons for this brand:

Business Description: {description}

Target Audience: {targeted_audience}

Author Bio: {author_bio}

Please suggest 3 different spokespersons who:
1. Would resonate with the target audience
2. Aligns with the brand's values and mission as described
3. Has the right personality and characteristics for this type of business
4. Would be credible and trustworthy for this industry

For each spokesperson, provide:
- Name (and indicate if they are real or fictional)
- Brief explanation (2-3 sentences) of why they would be perfect for this brand

Format your response as a JSON array with this structure:
[
  {{
    "name": "Person Name (Real/Fictional)",
    "explanation": "Brief explanation of why they would be perfect for this brand..."
  }},
  {{
    "name": "Person Name (Real/Fictional)", 
    "explanation": "Brief explanation of why they would be perfect for this brand..."
  }},
  {{
    "name": "Person Name (Real/Fictional)",
    "explanation": "Brief explanation of why they would be perfect for this brand..."
  }}
]

Make sure to provide diverse suggestions that would appeal to different aspects of the target audience.
"""

        # Use OpenAI to generate the brand spokesperson suggestions
        messages = [
            {"role": "system", "content": "You are a brand strategist who helps businesses choose the perfect spokesperson for their brand. Always respond with valid JSON format."},
            {"role": "user", "content": prompt}
        ]

        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                json={"model": "gpt-4.1", "messages": messages, "temperature": 0.7},
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
            )

            response_data = response.json()
            content = (
                response_data.get("choices", [{}])[0]
                .get("message", {})
                .get("content", "")
                .strip()
            )

            if not content:
                logger.error("OpenAI returned empty content")
                raise HTTPException(status_code=500, detail="Failed to generate brand spokesperson suggestions")
            
            logger.info(f"OpenAI response content length: {len(content)}")

            # Parse the JSON response
            try:
                import json
                suggestions = json.loads(content)
                
                # Validate the response format
                if not isinstance(suggestions, list) or len(suggestions) != 3:
                    raise ValueError("Invalid response format")
                
                for suggestion in suggestions:
                    if not isinstance(suggestion, dict) or "name" not in suggestion or "explanation" not in suggestion:
                        raise ValueError("Invalid suggestion format")
                
                logger.info(f"Successfully generated {len(suggestions)} brand spokesperson suggestions")
                return {"suggestions": suggestions}
                
            except (json.JSONDecodeError, ValueError) as e:
                logger.error(f"Failed to parse AI response as JSON: {e}")
                # Fallback: try to extract suggestions from text format
                fallback_suggestions = []
                lines = content.split('\n')
                current_suggestion = {}
                
                for line in lines:
                    line = line.strip()
                    if line.startswith('"name"') or line.startswith("name:"):
                        if current_suggestion:
                            fallback_suggestions.append(current_suggestion)
                        name = line.split(':', 1)[1].strip().strip('"').strip(',')
                        current_suggestion = {"name": name, "explanation": ""}
                    elif line.startswith('"explanation"') or line.startswith("explanation:"):
                        explanation = line.split(':', 1)[1].strip().strip('"').strip(',')
                        if current_suggestion:
                            current_suggestion["explanation"] = explanation
                
                if current_suggestion:
                    fallback_suggestions.append(current_suggestion)
                
                if len(fallback_suggestions) >= 3:
                    return {"suggestions": fallback_suggestions[:3]}
                else:
                    raise HTTPException(status_code=500, detail="Failed to parse brand spokesperson suggestions")

        except requests.exceptions.RequestException as e:
            logger.error(f"OpenAI API request error: {e}")
            raise HTTPException(
                status_code=500,
                detail="Failed to generate brand spokesperson due to API error.",
            )

    except Exception as e:
        logger.error(f"Error generating brand spokesperson: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

